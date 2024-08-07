from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
from .models import Document, Content
from django.contrib.auth.models import User 
from docx import Document as DocxDocument
import PyPDF2
from io import BytesIO
from django.http import HttpResponse
from django.conf import settings
import os
from .nlp_utils import improve_document_content


#  set max Size
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

@api_view(['POST'])
@permission_classes([IsAuthenticated])
def upload_document(request):
    user = request.user
    file = request.FILES.get('file')

    if not file:
        return Response({"error": "No file provided"}, status=status.HTTP_400_BAD_REQUEST)
    
    
    if file.size > MAX_FILE_SIZE:
        return Response({"error": "File size exceeds the limit of 10MB"}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Extract file name and extension
        file_name = file.name
        file_name, file_extension = os.path.splitext(file_name)
        
        # Create a new document entry with file name
        document = Document.objects.create(
            user=user,
            file_name=file.name 
        )
        
        # Extract content from file
        original_content = ''
        if file_extension == '.txt':
            original_content = file.read().decode('utf-8')
        elif file_extension == '.docx':
            doc = DocxDocument(file)
            for para in doc.paragraphs:
                original_content += para.text + '\n'
        elif file_extension == '.pdf':
            pdf = PyPDF2.PdfReader(file)
            for page in pdf.pages:
                original_content += page.extract_text() + '\n'
        else:
            return Response({"error": "Unsupported file format"}, status=status.HTTP_400_BAD_REQUEST)

        # Save 
        Content.objects.create(document=document, original_content=original_content)
        
        # Construct response data
        response_data = {
            "document_id": document.id,
            "file_name": document.file_name,
            "status": document.status,
            "upload_date": document.upload_date.isoformat(), 
            "content": original_content  
        }
        
        return Response(response_data, status=status.HTTP_201_CREATED)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_document(request, id):
    try:
        # Retrieve by ID
        document = Document.objects.get(id=id)
        # Retrieve the content associated with the document
        content = Content.objects.get(document=document)
        
        # response data
        response_data = {
            "document_id": document.id,
            "file_name": document.file_name,
            "status": document.status,
            "upload_date": document.upload_date.isoformat(),  
            "original_content": content.original_content,
            "improved_content": content.improved_content or ''  
        }
        
        return Response(response_data, status=status.HTTP_200_OK)
    except Document.DoesNotExist:
        return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
    except Content.DoesNotExist:
        return Response({"error": "Content not found for this document"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def improve_document(request, id):
    try:
       
        document = Document.objects.get(id=id)
        
        content = Content.objects.get(document=document)
        
        # NLP
        improved_content, suggestions = improve_document_content(content.original_content)
        
        # Save 
        content.improved_content = improved_content
        content.suggestions = suggestions
        content.save()
        
        
        response_data = {
            "document_id": document.id,
            "file_name": document.file_name,
            "status": document.status,
            "upload_date": document.upload_date.isoformat(), 
            "original_content": content.original_content,
            "improved_content": content.improved_content,
            "suggestions": suggestions  
        }
        
        return Response(response_data, status=status.HTTP_200_OK)
    except Document.DoesNotExist:
        return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
    except Content.DoesNotExist:
        return Response({"error": "Content not found for this document"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['PATCH'])
@permission_classes([IsAuthenticated])
def update_document_status(request, id):
    try:
        # Retrieve  by ID
        document = Document.objects.get(id=id)
        
        # Get the new status from the request data
        new_status = request.data.get('status')
        
        if not new_status:
            return Response({"error": "No status provided"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Update the document status
        document.status = new_status
        document.save()
        
        
        
        return Response({'message':"Document Status Updated Successfully"}, status=status.HTTP_200_OK)
    except Document.DoesNotExist:
        return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_all_documents(request):
    try:
        user = request.user
        # Retrieve all documents for the logged-in user
        documents = Document.objects.filter(user=user)
        
        response_data = []
        for document in documents:
            try:
                # Retrieve the content associated with the document
                content = Content.objects.get(document=document)
                doc_data = {
                    "document_id": document.id,
                    "file_name": document.file_name,
                    "status": document.status,
                    "upload_date": document.upload_date.isoformat(),  
                    "original_content": content.original_content,
                    "improved_content": content.improved_content or ''  
                }
                response_data.append(doc_data)
            except Content.DoesNotExist:
                continue
        
        return Response(response_data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
@permission_classes([IsAuthenticated])
def get_statistics(request):
    try:
        total_users = User.objects.count()
        total_uploaded_documents = Document.objects.count()
        total_approved_documents = Document.objects.filter(status='approved').count()
        total_rejected_documents = Document.objects.filter(status='rejected').count()
        total_documents = Document.objects.count()
        
        response_data = {
            'total_users': total_users,
            'total_uploaded_documents': total_uploaded_documents,
            'total_approved_documents': total_approved_documents,
            'total_rejected_documents': total_rejected_documents,
            'total_documents': total_documents
        }
        
        return Response(response_data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)



@api_view(['GET'])
@permission_classes([IsAuthenticated])
def generate_word_document(request, document_id):
    try:
        # Get the document and its improved content from the database
        doc = Document.objects.get(id=document_id)
        content = Content.objects.get(document=doc)
        
        # Get the improved content
        improved_content = content.improved_content or '' 
        
        # Define the organization name and document title
        organization_name = "AI DOCUMENT ASSISTANT"
        document_title = doc.file_name  

        # Load the template
        template_path = os.path.join(settings.BASE_DIR, 'templates', 'template.docx')
        
        if not os.path.exists(template_path):
            return Response({"error": "Template file not found"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        template = DocxDocument(template_path)

        # Replace placeholders in the template with actual content
        for paragraph in template.paragraphs:
            if '<<ORGANIZATION_NAME>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<ORGANIZATION_NAME>>', organization_name)
            if '<<DOCUMENT_TITLE>>' in paragraph.text:
                paragraph.text = paragraph.text.replace('<<DOCUMENT_TITLE>>', document_title)
            if '<<IMPROVED_CONTENT>>' in paragraph.text:
                # To avoid text replacement issues with formatting, handle in runs
                inline_shapes = paragraph.runs
                for run in inline_shapes:
                    if '<<IMPROVED_CONTENT>>' in run.text:
                        run.text = run.text.replace('<<IMPROVED_CONTENT>>', improved_content)

        # Save the document to a BytesIO object
        file_stream = BytesIO()
        template.save(file_stream)
        file_stream.seek(0)

        # Create the HTTP response
        response = HttpResponse(file_stream, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename={doc.file_name}_optimized.docx'

        return response

    except Document.DoesNotExist:
        return Response({"error": "Document not found"}, status=status.HTTP_404_NOT_FOUND)
    except Content.DoesNotExist:
        return Response({"error": "Content not found for this document"}, status=status.HTTP_404_NOT_FOUND)
    except FileNotFoundError:
        return Response({"error": "Template file not found"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


